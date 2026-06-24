"""Unit tests for document text extraction (PDF, DOCX, plain-text routing).

DOCX is generated for real with python-docx (a project dependency). PDF parsing
is exercised by stubbing ``pypdf.PdfReader`` so we test the extraction routing /
page-joining logic without bundling a binary PDF fixture.
"""
from __future__ import annotations

import io
import zipfile

import pytest

from src.activities.summarization import (
    DOCX_MIME,
    ODT_MIME,
    PDF_MIME,
    TXT_MIME,
    extract_text_from_bytes,
)


def test_extract_docx_real_document():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Bonjour le monde")
    doc.add_paragraph("Deuxième paragraphe")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text_from_bytes(buf.getvalue(), DOCX_MIME, "rapport.docx")
    assert "Bonjour le monde" in text
    assert "Deuxième paragraphe" in text


def test_extract_docx_by_filename_extension():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Contenu")
    buf = io.BytesIO()
    doc.save(buf)

    # mime omitted -> routed by the .docx extension
    text = extract_text_from_bytes(buf.getvalue(), "application/octet-stream", "x.docx")
    assert "Contenu" in text


def test_extract_pdf_joins_pages(monkeypatch):
    import pypdf

    class _FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class _FakeReader:
        def __init__(self, _stream):
            self.pages = [_FakePage("Page un"), _FakePage("Page deux")]

    monkeypatch.setattr(pypdf, "PdfReader", _FakeReader)

    text = extract_text_from_bytes(b"%PDF-1.4 fake", PDF_MIME, "doc.pdf")
    assert text == "Page un\nPage deux"


def test_extract_pdf_handles_none_page_text(monkeypatch):
    import pypdf

    class _FakePage:
        def extract_text(self):
            return None  # pypdf returns None for some pages

    class _FakeReader:
        def __init__(self, _stream):
            self.pages = [_FakePage()]

    monkeypatch.setattr(pypdf, "PdfReader", _FakeReader)

    assert extract_text_from_bytes(b"%PDF", PDF_MIME, "d.pdf") == ""


def test_extract_txt_utf8():
    assert extract_text_from_bytes("Été à Paris".encode("utf-8"), TXT_MIME, "n.txt") == "Été à Paris"


def test_extract_unsupported_mime_raises():
    with pytest.raises(ValueError):
        extract_text_from_bytes(b"data", "image/png", "x.png")


def test_extract_malformed_docx_raises():
    # A corrupt DOCX (not a valid zip) must raise rather than silently returning
    # empty text. The real python-docx parser is exercised here.
    with pytest.raises(Exception):
        extract_text_from_bytes(b"this is not a docx zip", DOCX_MIME, "broken.docx")


def test_extract_malformed_pdf_raises():
    # A corrupt PDF must raise rather than silently returning empty text. The real
    # pypdf parser is exercised here.
    with pytest.raises(Exception):
        extract_text_from_bytes(b"this is not a pdf", PDF_MIME, "broken.pdf")


# ---------------------------------------------------------------------------
# ODT extraction tests
# ---------------------------------------------------------------------------

def _build_odt(content_xml: str) -> bytes:
    """Return in-memory ODT bytes (a ZIP) containing the given content.xml."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_STORED) as zf:
        zf.writestr("content.xml", content_xml)
    return buf.getvalue()


_MINIMAL_CONTENT_XML = """\
<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
    xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">
  <office:body>
    <office:text>
      <text:p>Premier paragraphe</text:p>
      <text:h>Titre de section</text:h>
      <text:p>Deuxième paragraphe</text:p>
    </office:text>
  </office:body>
</office:document-content>
"""


def test_extract_odt_paragraph_and_heading_text():
    """Paragraphs and headings are extracted; boundaries are preserved as newlines."""
    data = _build_odt(_MINIMAL_CONTENT_XML)
    text = extract_text_from_bytes(data, ODT_MIME, "doc.odt")

    assert "Premier paragraphe" in text
    assert "Titre de section" in text
    assert "Deuxième paragraphe" in text
    # Boundaries: each block must be on its own line.
    lines = text.splitlines()
    assert "Premier paragraphe" in lines
    assert "Titre de section" in lines
    assert "Deuxième paragraphe" in lines


def test_extract_odt_by_filename_extension_fallback():
    """Generic MIME with an .odt filename routes to the ODT extractor."""
    data = _build_odt(_MINIMAL_CONTENT_XML)
    # application/octet-stream is a common generic fallback reported by some OSes.
    text = extract_text_from_bytes(data, "application/octet-stream", "rapport.odt")
    assert "Premier paragraphe" in text


def test_extract_odt_not_a_zip_raises():
    """Bytes that are not a valid ZIP raise ValueError (not silently return empty)."""
    with pytest.raises(ValueError, match="ODT file is not a valid ZIP"):
        extract_text_from_bytes(b"this is definitely not a zip", ODT_MIME, "bad.odt")


def test_extract_odt_missing_content_xml_raises():
    """A valid ZIP that lacks content.xml raises ValueError."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("styles.xml", "<root/>")
    with pytest.raises(ValueError, match="missing content.xml"):
        extract_text_from_bytes(buf.getvalue(), ODT_MIME, "no_content.odt")


def test_extract_odt_malformed_xml_raises():
    """A ZIP with broken XML in content.xml raises ValueError."""
    data = _build_odt("<<not valid xml>>")
    with pytest.raises(ValueError, match="not valid XML"):
        extract_text_from_bytes(data, ODT_MIME, "broken_xml.odt")
